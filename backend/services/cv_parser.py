"""
CV / Resume parser service.

Supports: PDF (.pdf), Word (.docx / .doc), plain text (.txt)

Parse strategy (no LLM required for basic extraction):
  1. extract_text_from_file()  → raw text from uploaded bytes
  2. parse_cv_local()          → rule-based extraction (always works, no tokens)
  3. parse_cv_llm()            → LLM enrichment (optional, skipped on 429 / no key)
  4. parse_cv()                → public API: local first, LLM if available
"""

import io
import json
import logging
import os
import re

logger = logging.getLogger(__name__)

EDUCATION_LEVELS = [
    "Primary school",
    "Junior high school (JHS)",
    "Senior high school (SHS) / Secondary",
    "TVET / Vocational certificate",
    "Bachelor degree or above",
    "No formal education",
]

# ── File text extraction ──────────────────────────────────────────────────────

def extract_text_from_file(content: bytes, filename: str) -> str:
    """Extract plain text from a CV file (PDF / DOCX / TXT)."""
    ext = os.path.splitext(filename.lower())[1]

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            return "\n".join(pages).strip()
        except ImportError:
            logger.warning("pdfplumber not installed — cannot parse PDF")
            return ""
        except Exception as e:
            logger.error("PDF extraction failed: %s", e)
            return ""

    if ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs).strip()
        except ImportError:
            logger.warning("python-docx not installed — cannot parse DOCX")
            return ""
        except Exception as e:
            logger.error("DOCX extraction failed: %s", e)
            return ""

    # TXT / fallback
    try:
        return content.decode("utf-8", errors="replace").strip()
    except Exception as e:
        logger.error("Text decode failed: %s", e)
        return ""


# ── Local rule-based parser (no LLM) ─────────────────────────────────────────

# Ordered highest → lowest; first match wins
_EDU_PATTERNS: list[tuple[str, list[str]]] = [
    ("Bachelor degree or above", [
        r"\b(ph\.?d|doctorate|doctor\s+of)\b",
        r"\bm\.?b\.?a\b",
        r"\b(master|m\.?sc|m\.?tech|m\.?e\.?|mca|m\.?eng)\b",
        r"\b(bachelor|b\.?sc|b\.?tech|b\.?e\.?|bca|b\.?eng|b\.?com|b\.?ed)\b",
        r"\b(degree|honours|hons|university|undergraduate|postgraduate)\b",
        r"\b(llb|mbbs|bds|md|be\b)",
    ]),
    ("TVET / Vocational certificate", [
        r"\b(diploma|vocational|polytechnic|tvet|bteb|nvq)\b",
        r"\biti\s*certificate\b",
        r"\b(certificate\s+course|professional\s+certificate)\b",
        r"\b(trade\s+certificate|apprentice)\b",
    ]),
    ("Senior high school (SHS) / Secondary", [
        r"\b(higher\s+secondary|senior\s+high|shs|hsc|a[\-\s]level)\b",
        r"\b(12th|class\s*12|grade\s*12|form\s*6)\b",
        r"\b(secondary\s+school|high\s+school|sec\s+school)\b",
    ]),
    ("Junior high school (JHS)", [
        r"\b(junior\s+high|middle\s+school|jhs|jsc|o[\-\s]level)\b",
        r"\b(10th|class\s*10|grade\s*10|form\s*4|ssc\b)\b",
        r"\b(basic\s+education|elementary\s+school)\b",
    ]),
    ("Primary school", [
        r"\b(primary\s+school|elementary|class\s*[1-6]|grade\s*[1-6])\b",
    ]),
]

# Section headers that signal experience content
_EXP_HEADERS = re.compile(
    r"(work\s+experience|employment|career\s+history|professional\s+experience"
    r"|experience|positions?|jobs?|internship)",
    re.IGNORECASE,
)

_SKIP_HEADERS = re.compile(
    r"(education|qualification|skill|language|reference|objective|summary|profile|award|certification|hobby)",
    re.IGNORECASE,
)


def _detect_education_level(text: str) -> str:
    lower = text.lower()
    for level, patterns in _EDU_PATTERNS:
        for pat in patterns:
            if re.search(pat, lower):
                return level
    return "Senior high school (SHS) / Secondary"


def _extract_experience_section(text: str) -> str:
    """
    Try to pull the experience section from the CV.
    Falls back to the whole text if no clear section headers are found.
    """
    lines = text.splitlines()
    in_exp = False
    exp_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if _EXP_HEADERS.search(stripped) and len(stripped) < 60:
            in_exp = True
            continue

        if in_exp and _SKIP_HEADERS.search(stripped) and len(stripped) < 60:
            in_exp = False
            continue

        if in_exp:
            exp_lines.append(stripped)

    if exp_lines:
        return " ".join(exp_lines)[:2000]

    # No labelled section — return whole text trimmed
    return text[:2000]


def parse_cv_local(cv_text: str) -> dict:
    """
    Rule-based CV parser. No LLM required.
    Returns {"education_level": str, "experience_text": str, "source": "local"}.
    """
    education_level = _detect_education_level(cv_text)
    experience_text = _extract_experience_section(cv_text).strip()

    # Truncate to a readable length
    if len(experience_text) > 800:
        experience_text = experience_text[:800].rsplit(" ", 1)[0] + "…"

    return {
        "education_level": education_level,
        "experience_text": experience_text,
        "source": "local",
    }


# ── LLM enrichment (optional) ────────────────────────────────────────────────

CV_EXTRACT_PROMPT = """You are a professional CV/resume parser.

Read the CV below and extract exactly two fields:

1. education_level — the highest level of education completed.
   Choose ONE of these exact strings:
   - "Primary school"
   - "Junior high school (JHS)"
   - "Senior high school (SHS) / Secondary"
   - "TVET / Vocational certificate"
   - "Bachelor degree or above"
   - "No formal education"

2. experience_text — a rich 3–5 sentence summary covering job titles, durations,
   key responsibilities, self-taught skills, projects and achievements.
   Be specific — include numbers, technologies, and outcomes.

CV content:
---
{cv_text}
---

Respond ONLY with this exact JSON (no markdown fences, no extra keys):
{{"education_level": "<level>", "experience_text": "<summary>"}}
"""


def _is_rate_limited_error(e: Exception) -> bool:
    msg = str(e).lower()
    return "429" in msg or "rate limit" in msg or "quota" in msg or "tokens per day" in msg


def parse_cv_llm(cv_text: str) -> dict | None:
    """
    Try to parse CV with LLM. Returns None on rate limit or any API error
    so the caller can fall back to the local parser result.
    """
    try:
        from backend.services.llm import generate_text
    except ImportError:
        return None

    truncated = cv_text[:3500]
    prompt    = CV_EXTRACT_PROMPT.format(cv_text=truncated)

    try:
        raw = generate_text(prompt).strip()

        # Strip markdown fences
        if raw.startswith("```"):
            parts = raw.split("```")
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith("json"):
                raw = raw[4:]

        result = json.loads(raw.strip())

        if result.get("education_level") not in EDUCATION_LEVELS:
            result["education_level"] = "Senior high school (SHS) / Secondary"

        result["source"] = "llm"
        return result

    except Exception as e:
        if _is_rate_limited_error(e):
            logger.warning("CV LLM skipped — rate limited (429). Using local parser.")
        else:
            logger.error("CV LLM parse failed: %s. Falling back to local parser.", e)
        return None


# ── Public API ────────────────────────────────────────────────────────────────

def parse_cv(cv_text: str) -> dict:
    """
    Parse a CV and return {education_level, experience_text}.

    Strategy:
      1. Always run local rule-based parser (fast, zero tokens).
      2. Try LLM enrichment — if it succeeds and the experience text is richer,
         use the LLM result; otherwise keep the local result.
      3. On any LLM error (including 429), silently use the local result.
    """
    if not cv_text.strip():
        return {"education_level": "", "experience_text": ""}

    local = parse_cv_local(cv_text)
    llm   = parse_cv_llm(cv_text)

    if llm and llm.get("experience_text") and len(llm["experience_text"]) > 30:
        logger.info("CV parsed via LLM (education=%s)", llm["education_level"])
        return {k: v for k, v in llm.items() if k != "source"}

    logger.info("CV parsed locally (education=%s)", local["education_level"])
    return {k: v for k, v in local.items() if k != "source"}
